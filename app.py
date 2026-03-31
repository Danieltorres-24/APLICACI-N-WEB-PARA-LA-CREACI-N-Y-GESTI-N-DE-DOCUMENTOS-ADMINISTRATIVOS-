from flask import Flask, render_template, request, jsonify, session, redirect, url_for, flash, send_file
from conexion import get_connection
import bcrypt
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import sqlite3
import os
from functools import wraps
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.oxml.ns import qn 
from docx.oxml import OxmlElement
from io import BytesIO
import logging
from pathlib import Path 
from docx.text.run import Run
from werkzeug.utils import secure_filename
from flask import current_app
from functools import wraps
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


app = Flask(__name__)
app.secret_key = 'CoDI'  

ADMIN_EMAIL = "oscardanielramireztorrez@gmail.com"
ADMIN_PASSWORD_HASH = generate_password_hash("123456")

def role_required(*roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'rol' not in session:
                return redirect(url_for('login'))
            if session['rol'] not in roles:
                return redirect(url_for('inicio'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator 

@app.route('/', methods=['GET'])
def root():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'IdUsuario' in session:
        if session.get('rol') == 'administrador':
            return redirect(url_for('administrador'))
        elif session.get('rol') == 'jefe':
            return redirect(url_for('jefe'))
        elif session.get('rol') == 'recepcion':
            return redirect(url_for('recepcion'))
        return redirect(url_for('inicio'))

    if request.method == 'POST':
        correo = request.form.get('correo')
        contrasena = request.form.get('contrasena')

        if not correo or not contrasena:
            return render_template('login.html', error='Por favor, complete todos los campos.')

        # Verificación para admin
        if correo == ADMIN_EMAIL:
            if check_password_hash(ADMIN_PASSWORD_HASH, contrasena):
                session.clear()
                session['rol'] = 'administrador'
                session['correo'] = ADMIN_EMAIL
                return redirect(url_for('administrador'))
            return render_template('login.html', error='Credenciales inválidas')

        # Verificación para usuarios normales
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        try:
            cursor.execute("""
                SELECT IdUsuario, nombre, correo, contrasena, cargo, siglas, area, rol 
                FROM usuarios 
                WHERE correo = %s
            """, (correo,))
            user = cursor.fetchone()

            if user:
                # Depuración: Verificar contraseña y rol
                app.logger.debug(f"Usuario encontrado: {user['nombre']}, Rol: {user['rol']}")
                
                if bcrypt.checkpw(contrasena.encode('utf-8'), user['contrasena'].encode('utf-8')):
                    session.clear()
                    session['IdUsuario'] = user['IdUsuario']
                    session['nombre'] = user['nombre']
                    session['correo'] = user['correo']
                    session['cargo'] = user['cargo']
                    session['siglas'] = user['siglas']
                    session['area'] = user['area']
                    session['rol'] = user['rol']

                    app.logger.debug(f"Login exitoso. Redirigiendo según rol: {user['rol']}")
                    
                    if user['rol'] == 'administrador':
                        return redirect(url_for('administrador'))
                    elif user['rol'] == 'jefe': 
                        return redirect(url_for('jefe'))
                    elif user['rol'] == 'recepcion':
                        return redirect(url_for('recepcion'))
                    elif user['rol'] == 'usuario' :
                        return redirect(url_for('inicio'))
                    return redirect(url_for('inicio'))
            
            return render_template('login.html', error='Credenciales inválidas')
        except Exception as e:
            app.logger.error(f"Error en login: {str(e)}")
            return render_template('login.html', error='Error al iniciar sesión')
        finally:
            cursor.close()
            conn.close()

    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear() 
    return redirect(url_for('login'))

@app.route('/index')
def inicio():
    if 'IdUsuario' not in session and 'rol' not in session:
        return redirect(url_for('login'))
    if 'IdUsuario' in session:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        try:
            # Datos del usuario
            cursor.execute("""
                SELECT nombre, correo, cargo, siglas, area
                FROM usuarios
                WHERE IdUsuario = %s
            """, (session['IdUsuario'],))
            user_data = cursor.fetchone()
            
            if not user_data:
                return redirect(url_for('logout'))

            # Guardar en sesión
            session['nombre'] = user_data['nombre']
            session['correo'] = user_data['correo']
            session['cargo'] = user_data['cargo']
            session['siglas'] = user_data['siglas']
            session['area'] = user_data['area']

            # Ahora traemos los folios que creó el usuario
            cursor.execute("""
                SELECT IdSalida, origenArea, folio, tipo, fecha_registro
                FROM salidas
                ORDER BY IdSalida DESC
                LIMIT 5
            """)
            folios = cursor.fetchall()
            return render_template('index.html', folios=folios)
        except Exception as e:
            app.logger.error(f"Error al cargar index: {str(e)}")
            return redirect(url_for('logout'))
        finally:
            cursor.close()
            conn.close()
    
    return redirect(url_for('administrador'))


@app.route('/formatos')
def formatos():
    if 'IdUsuario' not in session or 'rol' not in session:
        return redirect(url_for('login'))
    return render_template('formatos.html')


@app.route('/copiar_doc', methods=['GET'])
def copiar_doc():
    try:
        template_path = os.path.join(current_app.root_path, "static", "etiquetas", "Formatodeoficio.docx")
        if not os.path.exists(template_path):
            return jsonify({'error': 'El documento no existe'}), 404

        doc = Document(template_path)

        contenido = []
        for p in doc.paragraphs:
            if p.text.strip():
                contenido.append(p.text.strip())

        texto = "\n".join(contenido)
        print("Ruta del documento:", template_path)


        return jsonify({"texto": texto})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/subir_plantilla', methods=['POST'])
def subir_plantilla():
    if 'IdUsuario' not in session or 'rol' not in session:
        flash('No autorizado', 'error')
        return redirect(url_for('formatos'))

    if 'documento' not in request.files:
        flash('No se envió archivo', 'error')
        return redirect(url_for('formatos'))

    archivo = request.files['documento']
    if archivo.filename == '' or not archivo.filename.lower().endswith('.docx'):
        flash('Archivo no válido. Solo se permiten archivos .docx', 'error')
        return redirect(url_for('formatos'))

    carpeta_destino = os.path.join(current_app.root_path, "static", "documentos")
    os.makedirs(carpeta_destino, exist_ok=True)

    nombre_archivo = f"formato_{session['IdUsuario']}.docx"
    destino = os.path.join(carpeta_destino, nombre_archivo)

    try:
        archivo.save(destino)
        return jsonify({'mensaje': f'Plantilla guardada como {nombre_archivo}'})
    except Exception as e:
        flash(f'Error al guardar plantilla: {str(e)}', 'error')

    return redirect(url_for('formatos'))



@app.route('/admin', methods=['GET'])
def administrador():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM usuarios")
    usuarios = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template('admin.html', usuarios=usuarios)

@app.route('/jefe')
def jefe():
    if 'IdUsuario' not in session or 'rol' not in session:
        return redirect(url_for('login'))

    if session.get('rol') != 'jefe':
        return redirect(url_for('inicio'))

    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        # Datos del jefe
        cursor.execute("""
            SELECT nombre, correo, cargo, siglas, area
            FROM usuarios
            WHERE IdUsuario = %s
        """, (session['IdUsuario'],))
        jefe_data = cursor.fetchone()

        if not jefe_data:
            return redirect(url_for('logout'))

        session['nombre'] = jefe_data['nombre']
        session['correo'] = jefe_data['correo']
        session['cargo'] = jefe_data['cargo']
        session['siglas'] = jefe_data['siglas']
        session['area'] = jefe_data['area']

        # Folios de todo su equipo / área
        cursor.execute("""
            SELECT s.IdSalida, s.origenArea, s.folio, s.tipo, s.fecha_registro, u.nombre AS usuario
            FROM salidas s
            JOIN usuarios u ON s.IdUsuario = u.IdUsuario
            WHERE s.origenArea = %s
            ORDER BY s.IdSalida DESC
            LIMIT 5
        """, (session['area'],))
        folios = cursor.fetchall()

        return render_template('jefe.html', folios=folios)

    except Exception as e:
        app.logger.error(f"Error al cargar jefe: {str(e)}")
        return redirect(url_for('logout'))
    finally:
        cursor.close()
        conn.close()

@app.route('/recepcion')
@role_required('recepcion', 'administrador')
def recepcion():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT nombre, correo, cargo, siglas, area
            FROM usuarios
            WHERE IdUsuario = %s
        """, (session['IdUsuario'],))
        user_data = cursor.fetchone()

        cursor.execute("""
            SELECT IdEntrada, folio, area_origen, resumen, fecha, documento_oficio, documento_anexo, cantidad_anexos
            FROM entrada
            ORDER BY fecha DESC
        """)
        registros = cursor.fetchall()

        # Mapear nombres para JS
        registros_js = [{
            "IdEntrada": r["IdEntrada"],
            "folio": r["folio"],
            "area": r["area_origen"],
            "descripcion": r["resumen"],
            "documento_oficio": r["documento_oficio"],
            "documento_anexo": r["documento_anexo"],
            "cantidad_anexos": r["cantidad_anexos"],
            "fecha": r["fecha"].strftime("%Y-%m-%d %H:%M:%S")
        } for r in registros]

        return render_template('recepcion.html', user=user_data, folios=folios, registros=registros_js)
    finally:
        cursor.close()
        conn.close()


BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'static', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


@app.route("/upload", methods=["POST"])
def upload_file():
    if "documento_oficio" in request.files:
        file = request.files["documento_oficio"]
    elif "documento_anexo" in request.files:
        file = request.files["documento_anexo"]
    else:
        return jsonify({"error": "No se recibió archivo"}), 400

    if file and file.filename.endswith(".pdf"):
        filename = secure_filename(file.filename)
        filepath = os.path.join("static/uploads", filename)
        file.save(filepath)
        return jsonify({"filename": filename})
    return jsonify({"error": "Formato no permitido"}), 400



ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



# ------------------ Obtener registros ------------------
@app.route("/api/entrada", methods=["GET"])
def get_entradas():
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT IdEntrada, folio, area_origen, resumen, fecha,
               documento_oficio, documento_anexo, cantidad_anexos
        FROM entrada
        ORDER BY fecha DESC
    """)
    registros = cursor.fetchall()
    cursor.close()
    conn.close()

    registros_js = [{
        "IdEntrada": r["IdEntrada"],
        "folio": r["folio"],
        "area": r["area_origen"],
        "descripcion": r["resumen"],
        "documento_oficio": r["documento_oficio"],
        "documento_anexo": r["documento_anexo"],
        "cantidad_anexos": r["cantidad_anexos"],
        "fecha": r["fecha"].strftime("%Y-%m-%d %H:%M:%S")
    } for r in registros]

    return jsonify(registros_js)

# ------------------ Crear registro ------------------
@app.route("/api/entrada", methods=["POST"])
def create_entrada():
    data = request.get_json()
    cantidad_anexos = data.get("cantidad_anexos", "0")
    if not cantidad_anexos.isdigit():
        cantidad_anexos = "0"

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO entrada (folio, area_origen, resumen,
                             documento_oficio, documento_anexo,
                             cantidad_anexos, fecha)
        VALUES (%s, %s, %s, %s, %s, %s, NOW())
    """, (
        data["folio"], data["area"], data["descripcion"],
        data["documento_oficio"], data["documento_anexo"],
        cantidad_anexos
    ))
    conn.commit()
    cursor.close()
    conn.close()
    return jsonify({"success": True}), 201

# ------------------ Actualizar registro ------------------
@app.route("/api/entrada/<int:idEntrada>", methods=["PUT"])
def update_entrada(idEntrada):
    data = request.get_json()
    cantidad_anexos = data.get("cantidad_anexos", "0")
    if not cantidad_anexos.isdigit():
        cantidad_anexos = "0"

    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""
        UPDATE entrada
        SET folio=%s, area_origen=%s, resumen=%s,
            documento_oficio=%s, documento_anexo=%s,
            cantidad_anexos=%s
        WHERE IdEntrada=%s
    """, (
        data["folio"], data["area"], data["descripcion"],
        data["documento_oficio"], data["documento_anexo"],
        cantidad_anexos,
        idEntrada
    ))
    conn.commit()
    cursor.close()
    conn.close()
    return jsonify({"success": True})

# ------------------ Borrar registro ------------------
@app.route("/api/entrada/<int:idEntrada>", methods=["DELETE"])
def delete_entrada(idEntrada):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM entrada WHERE IdEntrada=%s", (idEntrada,))
    conn.commit()
    cursor.close()
    conn.close()
    return jsonify({"success": True})
















@app.route('/add_user', methods=['POST'])
def add_user():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()

    try:
        hashed_password = bcrypt.hashpw(data['contrasena'].encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

        cursor.execute("""
            INSERT INTO usuarios (nombre, correo, contrasena, cargo, siglas, area, rol)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (data['nombre'], data['correo'], hashed_password, data['cargo'], data['siglas'], data['area'], data['rol']))
        
        new_id = cursor.lastrowid 
        conn.commit()
        return jsonify({"message": "Usuario agregado exitosamente.", "id": new_id}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_user/<int:id>', methods=['POST'])
def edit_user(id):
    conn = get_connection()
    cursor = conn.cursor()

    try:
        data = request.get_json()
        if 'contrasena' in data and data['contrasena']:
            hashed_password = bcrypt.hashpw(data['contrasena'].encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            cursor.execute("""
                UPDATE usuarios
                SET nombre = %s, correo = %s, contrasena = %s, cargo = %s, siglas = %s, area = %s, rol = %s
                WHERE Idusuario = %s
            """, (data['nombre'], data['correo'], hashed_password, data['cargo'], data['siglas'], data['area'], data['rol'], id))
        else:
            cursor.execute("""
                UPDATE usuarios
                SET nombre = %s, correo = %s, cargo = %s, siglas = %s, area = %s, rol = %s
                WHERE Idusuario = %s
            """, (data['nombre'], data['correo'], data['cargo'], data['siglas'], data['area'], data['rol'], id))

        conn.commit()
        return jsonify({"message": "Usuario actualizado exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_user/<int:id>', methods=['POST'])
def delete_user(id):
       print(f"Eliminando usuario con ID: {id}")  
       conn = get_connection()
       cursor = conn.cursor()

       try:
           cursor.execute("DELETE FROM usuarios WHERE Idusuario = %s", (id,))
           conn.commit()
           return jsonify({"message": "Usuario eliminado exitosamente."}), 204
       except Exception as e:
           print(f"Error al eliminar: {str(e)}")  
           return jsonify({"error": str(e)}), 400
       finally:
           cursor.close()
           conn.close()
   
@app.route('/add_unidad', methods=['POST'])
def add_unidad():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO unidades (nombre, descripcion, localidad)
            VALUES (%s, %s, %s)
        """, (data['nombre'], data['descripcion'], data['localidad']))
        conn.commit()
        
        return jsonify({"message": "Unidad agregada exitosamente."}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_unidad/<int:id>', methods=['POST'])
def edit_unidad(id):
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE unidades
            SET nombre = %s, descripcion = %s, localidad = %s
            WHERE Idunidad = %s
        """, (data['nombre'], data['descripcion'], data['localidad'], id))
        conn.commit()
        
        return jsonify({"message": "Unidad actualizada exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_unidad/<int:id>', methods=['POST'])
def delete_unidad(id):
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM unidades WHERE Idunidad = %s", (id,))
        conn.commit()
        return jsonify({"message": "Unidad eliminada exitosamente."}), 204
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/get_unidades', methods=['GET'])
def get_unidades():
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT Idunidad, nombre, descripcion, localidad, fecha FROM unidades")
        unidades = cursor.fetchall()
        
        result = []
        for unidad in unidades:
            result.append({
                "id": unidad[0],
                "nombre": unidad[1],
                "descripcion": unidad[2],
                "localidad": unidad[3],
                "fecha": unidad[4].isoformat() 
            })
        
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()
        
@app.route('/get_areas', methods=['GET'])
def get_areas():
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT IdArea, nombre, UI, CC, siglas, fecha FROM areas")
        areas = cursor.fetchall()
        
        result = []
        for area in areas:
            result.append({
                "id": area[0],
                "nombre": area[1],
                "UI": area[2],
                "CC": area[3],
                "siglas": area[4],
                "fecha": area[5]
            })
        
        return jsonify(result), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/add_area', methods=['POST'])
def add_area():
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO areas (nombre, UI, CC, siglas, fecha)
            VALUES (%s, %s, %s, %s, %s)
        """, (data['nombre'], data['UI'], data['CC'], data['siglas'], data['fecha']))
        conn.commit()
        
        return jsonify({"message": "Área agregada exitosamente."}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/edit_area/<int:id>', methods=['POST'])
def edit_area(id):
    data = request.get_json()
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE areas
            SET nombre = %s, UI = %s, CC = %s, siglas = %s, fecha = %s
            WHERE IdArea = %s
        """, (data['nombre'], data['UI'], data['CC'], data['siglas'], data['fecha'], id))
        conn.commit()
        
        return jsonify({"message": "Área actualizada exitosamente."}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/delete_area/<int:id>', methods=['POST'])
def delete_area(id):
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM areas WHERE IdArea = %s", (id,))
        conn.commit()
        return jsonify({"message": "Área eliminada exitosamente."}), 204
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    finally:
        cursor.close()
        conn.close()

@app.route('/destinatarios')
def mostrar_destinatarios():
    return render_template('destinatarios.html')

@app.route('/api/destinatarios', methods=['GET', 'POST'])
def manejar_destinatarios():
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        if request.method == 'GET':
            page = request.args.get('page', 1, type=int)
            per_page = 5
            search = request.args.get('search', '', type=str)
            
            query = 'SELECT IdDestino AS id, nombre, cargo FROM destinatarios'
            params = []
            
            if search:
                query += ' WHERE nombre LIKE %s OR cargo LIKE %s'
                params.extend([f'%{search}%', f'%{search}%'])
            
            query += ' LIMIT %s OFFSET %s'
            params.extend([per_page, (page - 1) * per_page])
            
            cursor.execute(query, params)
            destinatarios = cursor.fetchall()
            
            count_query = 'SELECT COUNT(*) AS total FROM destinatarios'
            if search:
                count_query += ' WHERE nombre LIKE %s OR cargo LIKE %s'
                cursor.execute(count_query, [f'%{search}%', f'%{search}%'])
            else:
                cursor.execute(count_query)
                
            total_count = cursor.fetchone()['total']
            total_pages = (total_count + per_page - 1) // per_page
            
            return jsonify({
                'destinatarios': destinatarios,
                'total_pages': total_pages,
                'current_page': page
            })
            
        elif request.method == 'POST':
            data = request.get_json()
            nombre = data.get('nombre')
            cargo = data.get('cargo')
            
            if not nombre or not cargo:
                return jsonify({'error': 'Nombre y cargo son requeridos'}), 400
            
            cursor.execute(
                'INSERT INTO destinatarios (nombre, cargo) VALUES (%s, %s)',
                (nombre, cargo)
            )
            destinatario_id = cursor.lastrowid
            conn.commit()
            
            cursor.execute(
                'SELECT IdDestino AS id, nombre, cargo FROM destinatarios WHERE IdDestino = %s', 
                (destinatario_id,)
            )
            destinatario = cursor.fetchone()
            
            return jsonify(destinatario), 201
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@app.route('/api/destinatarios/<int:id>', methods=['PUT', 'DELETE'])
def manejar_destinatario(id):
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)
        
        if request.method == 'PUT':
            data = request.get_json()
            nombre = data.get('nombre')
            cargo = data.get('cargo')
            
            if not nombre or not cargo:
                return jsonify({'error': 'Nombre y cargo son requeridos'}), 400
            
            cursor.execute(
                'UPDATE destinatarios SET nombre = %s, cargo = %s WHERE IdDestino = %s',
                (nombre, cargo, id)
            )
            
            if cursor.rowcount == 0:
                return jsonify({'error': 'Destinatario no encontrado'}), 404
            
            conn.commit()
            
            cursor.execute(
                'SELECT IdDestino AS id, nombre, cargo FROM destinatarios WHERE IdDestino = %s', 
                (id,)
            )
            destinatario = cursor.fetchone()
            
            return jsonify(destinatario)
            
        elif request.method == 'DELETE':
            cursor.execute('DELETE FROM destinatarios WHERE IdDestino = %s', (id,))
            
            if cursor.rowcount == 0:
                return jsonify({'error': 'Destinatario no encontrado'}), 404
            
            conn.commit()
            return jsonify({'message': 'Destinatario eliminado correctamente'}), 200
            
    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()
        

@app.route('/folios')
def folios():
    """Renderiza la página de folios con las áreas disponibles"""
    conn = get_connection()
    if not conn:
        return "Error de conexión a la base de datos", 500

    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT IdArea, UI, CC, nombre FROM areas ORDER BY nombre ASC")
        areas = cursor.fetchall()
    except Exception as e:
        return f"Error: {e}", 500
    finally:
        cursor.close()
        conn.close()

    return render_template('folios.html', areas=areas)


@app.route('/api/folios', methods=['GET', 'POST'])
def manejar_folios():
    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión a la base de datos'}), 500

    try:
        cursor = conn.cursor(dictionary=True)

        # =========================
        # LISTAR FOLIOS
        # =========================
        if request.method == 'GET':
            if 'IdUsuario' not in session:
                return jsonify({'error': 'No autorizado'}), 401

            user_id = session['IdUsuario']
            user_role = session.get('rol', 'usuario')
            user_area = session.get('area')

            page = request.args.get('page', 1, type=int)
            per_page = request.args.get('per_page', 5, type=int)
            search = request.args.get('search', '', type=str)

            query = """
                SELECT s.IdSalida, s.folio, s.origenArea, s.tipo, s.fecha_registro, s.Asunto, s.acuse
                FROM salidas s
                JOIN usuarios u ON s.IdUsuario = u.IdUsuario
                WHERE 1=1
            """
            params = []

            if user_role == 'usuario':
                query += " AND s.IdUsuario = %s"
                params.append(user_id)
            elif user_role in ('recepcion', 'jefe'):
                query += " AND u.area = %s"
                params.append(user_area)

            if search:
                query += " AND (s.folio LIKE %s OR s.origenArea LIKE %s OR s.Asunto LIKE %s OR s.tipo LIKE %s)"
                params.extend([f"%{search}%", f"%{search}%", f"%{search}%", f"%{search}%"])


            query += " ORDER BY s.fecha_registro DESC, s.IdSalida DESC LIMIT %s OFFSET %s"
            params.extend([per_page, (page - 1) * per_page])
            cursor.execute(query, params)
            folios = cursor.fetchall()

            # Total registros
            count_query = "SELECT COUNT(*) AS total FROM salidas s JOIN usuarios u ON s.IdUsuario = u.IdUsuario WHERE 1=1"
            count_params = []
            if user_role == 'usuario':
                count_query += " AND s.IdUsuario = %s"
                count_params.append(user_id)
            elif user_role in ('recepcion', 'jefe'):
                count_query += " AND u.area = %s"
                count_params.append(user_area)

            if search:
                count_query += " AND (s.folio LIKE %s OR s.origenArea LIKE %s)"
                count_params.extend([f"%{search}%", f"%{search}%"])

            cursor.execute(count_query, count_params)
            total_count = cursor.fetchone()['total']
            total_pages = (total_count + per_page - 1) // per_page

            return jsonify({'folios': folios, 'totalPages': total_pages})

        # =========================
        # CREAR NUEVO FOLIO
        # =========================
        elif request.method == 'POST':
            if 'IdUsuario' not in session:
                return jsonify({'error': 'No autorizado'}), 401

            tipo = request.form.get('tipo')
            asunto = request.form.get('asunto').strip()

            if tipo not in ('memorándum', 'oficio'):
                return jsonify({'error': 'Tipo inválido'}), 400

            area_nombre = session.get('area')
            if not area_nombre:
                return jsonify({'error': 'Área no encontrada en la sesión'}), 400

            user_id = session['IdUsuario']
            user_siglas = session.get('siglas', 'DS')
            year = datetime.now().year

            cursor.execute("SELECT IdArea, UI, CC, nombre FROM areas WHERE nombre = %s", (area_nombre,))
            area = cursor.fetchone()
            if not area:
                return jsonify({'error': 'Área no encontrada en la base de datos'}), 404

            cursor.execute("""
                SELECT MAX(CAST(SUBSTRING_INDEX(SUBSTRING_INDEX(folio, '/', -2), '/', 1) AS UNSIGNED)) AS max_num
                FROM salidas
                WHERE tipo = %s AND YEAR(fecha_registro) = %s AND origenArea = %s
            """, (tipo, year, area['nombre']))
            result = cursor.fetchone()
            next_num = (result['max_num'] or 0) + 1

            folio_generado = f"{area['UI']}/{area['CC']}/{user_siglas}/{str(next_num).zfill(4)}/{year}"

            cursor.execute("""
                INSERT INTO salidas (origenArea, Asunto, folio, tipo, IdUsuario, fecha_registro)
                VALUES (%s, %s, %s, %s, %s, NOW())
            """, (area['nombre'], asunto, folio_generado, tipo, user_id))
            conn.commit()

            folio_id = cursor.lastrowid
            cursor.execute("""
                SELECT IdSalida, folio, Asunto, origenArea, tipo, fecha_registro
                FROM salidas WHERE IdSalida = %s
            """, (folio_id,))
            nuevo_folio = cursor.fetchone()

            return jsonify(nuevo_folio), 201

    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/folios/<int:folio_id>', methods=['DELETE'])
def eliminar_folio(folio_id):
    if 'IdUsuario' not in session:
        return jsonify({'error': 'No autorizado'}), 401

    conn = get_connection()
    if not conn:
        return jsonify({'error': 'Error de conexión'}), 500

    try:
        cursor = conn.cursor()
        cursor.execute("SELECT IdSalida FROM salidas WHERE IdSalida = %s", (folio_id,))
        folio = cursor.fetchone()
        if not folio:
            return jsonify({'error': 'Folio no encontrado'}), 404

        cursor.execute("DELETE FROM salidas WHERE IdSalida = %s", (folio_id,))
        conn.commit()

        return jsonify({'message': 'Folio eliminado correctamente'}), 200

    except Exception as e:
        conn.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

UPLOAD_FOLDER = 'static/acuse'
ALLOWED_EXTENSIONS = {'pdf'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

import os
from flask import request, jsonify, current_app

@app.route('/api/folios/acuse/<int:folio_id>', methods=['POST'])
def subir_acuse(folio_id):
    if 'IdUsuario' not in session:
        return jsonify({'error': 'No autorizado'}), 401

    if 'acuse' not in request.files:
        return jsonify({'error': 'No se envió archivo'}), 400

    file = request.files['acuse']
    if file.filename == '':
        return jsonify({'error': 'Archivo vacío'}), 400

    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Solo se permiten archivos PDF'}), 400

    try:
        # Crear carpeta si no existe
        upload_folder = os.path.join(current_app.root_path, 'static', 'acuse')
        os.makedirs(upload_folder, exist_ok=True)

        # Nombre del archivo: acuse_folioId.pdf
        filename = f"acuse_{folio_id}.pdf"
        file_path = os.path.join(upload_folder, filename)
        file.save(file_path)

        # Guardar ruta en base de datos
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("UPDATE salidas SET acuse = %s WHERE IdSalida = %s", (filename, folio_id))
        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({'message': 'Acuse subido correctamente', 'file': filename}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@app.route('/documento')
def editar_documento():
    folio_id = request.args.get('folio_id', type=int)
    if not folio_id:
        flash('No se especificó folio', 'error')
        return redirect(url_for('folios'))

    conn = get_connection()
    if not conn:
        flash('Error de conexión a la base de datos', 'error')
        return redirect(url_for('folios'))

    try:
        cursor = conn.cursor(dictionary=True)

        # =========================
        # Datos del folio y usuario
        # =========================
        cursor.execute("""
            SELECT s.*, u.nombre AS usuario_nombre, u.IdArea
            FROM salidas s
            JOIN usuarios u ON s.IdUsuario = u.IdUsuario
            WHERE s.IdSalida = %s
        """, (folio_id,))
        folio = cursor.fetchone()
        if not folio:
            flash('Folio no encontrado', 'error')
            return redirect(url_for('folios'))


        # =========================
        # Áreas disponibles
        # =========================
        cursor.execute("SELECT IdArea AS IdArea, nombre, UI, CC FROM areas ORDER BY nombre")
        areas = cursor.fetchall()

        # =========================
        # Destinatarios disponibles
        # =========================
        cursor.execute("SELECT IdDestino AS id, nombre, cargo FROM destinatarios ORDER BY nombre")
        destinatarios = cursor.fetchall()

        # =========================
        # Unidades disponibles
        # =========================
        cursor.execute("SELECT IdUnidad AS id, nombre FROM unidades ORDER BY nombre")
        unidades = cursor.fetchall()

        return render_template(
            'documento.html',
            folio=folio,
            folio_id=folio_id,
            areas=areas,
            destinatarios=destinatarios,
            unidades=unidades
        )

    except Exception as e:
        print(f"ERROR en /documento: {str(e)}")
        flash('Error al cargar el editor de documento', 'error')
        return redirect(url_for('folios'))
    finally:
        cursor.close()
        conn.close()

def insertar_html_en_docx(doc, html_cuerpo, marcador="{{CUERPO}}"):
    soup = BeautifulSoup(html_cuerpo, 'html.parser')
    for p in doc.paragraphs:
        if marcador in p.text:
            parent = p._element
            # Limpiar el párrafo original
            for run in p.runs:
                run.text = ""

            # Insertar cada elemento <p> o <table>
            for elem in soup.contents:
                if elem.name == 'p':
                    new_p = OxmlElement('w:p')
                    parent.addprevious(new_p)
                    paragraph = Paragraph(new_p, doc)
                    for sub in elem.contents:
                        if isinstance(sub, str):
                            paragraph.add_run(sub)
                        elif sub.name == 'strong':
                            paragraph.add_run(sub.get_text()).bold = True
                        elif sub.name == 'em':
                            paragraph.add_run(sub.get_text()).italic = True
                        elif sub.name == 'u':
                            paragraph.add_run(sub.get_text()).underline = True
                
                elif elem.name == 'table':
                    rows = elem.find_all('tr')
                    if not rows: 
                        continue
                    cols = rows[0].find_all(['td','th'])
                    table = doc.add_table(rows=len(rows), cols=len(cols))
                    table.style = 'Table Grid'
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td','th'])
                        for j, cell in enumerate(cells):
                            text = ''.join(cell.strings).strip()
                            table.cell(i,j).text = text

            # Borrar marcador original
            p.clear()
            break


@app.route('/generar_documento', methods=['POST'])
def generar_documento():
    conn = None
    cursor = None
    try:
        # 1) Sesión y JSON
        if 'IdUsuario' not in session:
            return jsonify({'error': 'No autorizado'}), 401
        if not request.is_json:
            return jsonify({'error': 'Se esperaba formato JSON'}), 400

        data = request.get_json()

        # 2) Validaciones
        required_fields = ['folio_id', 'destinatarios', 'cuerpo', 'elaborador', 'autorizo', 'copia']
        for campo in required_fields:
            if campo not in data:
                return jsonify({'error': f'Campo {campo} es requerido'}), 400

        if not isinstance(data['destinatarios'], list) or len(data['destinatarios']) == 0:
            return jsonify({'error': 'Se requiere al menos un destinatario'}), 400

        # 3) Obtener folio y área desde DB
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT folio, origenArea
            FROM salidas
            WHERE IdSalida = %s
        """, (data['folio_id'],))
        folio_db = cursor.fetchone()
        if not folio_db:
            return jsonify({'error': 'Folio no encontrado'}), 404

        numero_oficio = folio_db['folio']
        area_nombre = folio_db['origenArea'] or ''

        # 4) Fecha en español
        meses = ['enero','febrero','marzo','abril','mayo','junio','julio',
                 'agosto','septiembre','octubre','noviembre','diciembre']
        now = datetime.now()
        fecha_formateada = f"Acapulco de Juárez, Gro., a {now.day} de {meses[now.month-1]} de {now.year}"

        # 5) Texto de destinatarios
        destinatarios_texto = "\n".join(
            f"{(d.get('nombre') or '').strip()}\n{(d.get('cargo') or '').strip()}".strip()
            for d in data['destinatarios']
            if (d.get('nombre') or '').strip()
        ).strip()
        if not destinatarios_texto:
            return jsonify({'error': 'Se requiere al menos un destinatario con nombre'}), 400

        # 6) Plantilla del usuario
        plantilla_path = os.path.join(
            current_app.root_path, "static", "documentos",
            f"formato_{session['IdUsuario']}.docx"
        )
        if not os.path.exists(plantilla_path):
            return jsonify({'error': 'Plantilla de documento no encontrada'}), 500

        doc = Document(plantilla_path)

        # 7) Insertar cuerpo HTML en la posición de {{CUERPO}}
        html_cuerpo = data.get('cuerpo', '')
        insertar_html_en_docx(doc, html_cuerpo)

        # 8) Reemplazos restantes (incluyendo AUTORIZO y COPIA)
        replacements = {
            "{{UNIDAD}}": "Unidad de medicina familiar No9",
            "{{AREA}}": area_nombre,
            "{{NUMERO_OFICIO}}": f"Of. N° {numero_oficio}",
            "{{FECHA}}": fecha_formateada,
            "{{DESTINATARIOS}}": destinatarios_texto,
            "{{ELABORADOR}}": data.get('elaborador', ''),
            "{{AUTORIZO}}": data.get('autorizo', ''),
            "{{COPIA}}": data.get('copia', '')
        }

        def replace_in_paragraph(paragraph, replacements):
            if not paragraph.runs:
                return
            full_text = "".join(run.text for run in paragraph.runs)
            changed = False
            for search, repl in replacements.items():
                if search in full_text:
                    full_text = full_text.replace(search, repl)
                    changed = True
            if changed:
                for run in paragraph.runs:
                    run.text = ""
                paragraph.runs[0].text = full_text

        def replace_in_docx(doc, replacements):
            for p in doc.paragraphs:
                replace_in_paragraph(p, replacements)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p, replacements)

        replace_in_docx(doc, replacements)

        # 9) Enviar archivo
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        nombre_archivo = f"Oficio_{numero_oficio.replace('/', '_')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nombre_archivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logging.error(f'Error al generar documento: {str(e)}', exc_info=True)
        return jsonify({'error': f'Error al generar el documento: {str(e)}'}), 500

    finally:
        if cursor:
            cursor.close()
        if conn and conn.is_connected():
            conn.close()



if __name__ == '__main__':
    #app.run(debug=True, host='0.0.0.0', port=5000)
    #app.run(debug=True, port=5000)
    app.run(debug=True, host='localhost', port=5000, ssl_context=('localhost.pem', 'localhost-key.pem'))
    #app.run(debug=True, host='192.168.1.154', port=5000, ssl_context=('192.168.1.154.pem', '192.168.1.154-key.pem'))
    #app.run(debug=True, host='192.168.1.92', port=5000, ssl_context=('192.168.1.92.pem', '192.168.1.92-key.pem'))